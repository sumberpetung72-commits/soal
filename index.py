import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO

# --- KONFIGURASI API & BRANDING ---
# Ganti dengan API Key Anda atau gunakan st.secrets untuk keamanan
API_KEY = "ISI_API_KEY_ANDA_DI_SINI" 
NAMA_APLIKASI = "GuruAI Pro"
NAMA_PEMBUAT = "Nama Anda" # GANTI DENGAN NAMA ANDA

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# --- FUNGSI PENDUKUNG ---
def buat_file_word(konten):
    doc = Document()
    doc.add_heading(f'Hasil Perangkat Ujian - {NAMA_APLIKASI}', 0)
    doc.add_paragraph(konten)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- TAMPILAN UI (CSS) ---
st.set_page_config(page_title=NAMA_APLIKASI, page_icon="🎓", layout="wide")

st.markdown(f"""
    <style>
    .stApp {{ background-color: #fcfcfc; }}
    .footer {{
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: white;
        color: #666;
        text-align: center;
        padding: 10px;
        border-top: 1px solid #eee;
        font-size: 14px;
    }}
    .main-card {{
        background-color: white;
        padding: 20px;
        border-radius: 15px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }}
    </style>
    """, unsafe_allow_html=True)

# --- SIDEBAR (PANDUAN) ---
with st.sidebar:
    st.title(f"🎓 {NAMA_APLIKASI}")
    st.success(f"Dibuat oleh: **{NAMA_PEMBUAT}**")
    st.divider()
    st.header("📖 Panduan Cepat")
    st.write("""
    1. **Pilih Metode:** Masukkan teks manual atau Upload file PDF materi.
    2. **Atur Parameter:** Tentukan jenjang dan jumlah soal.
    3. **Generate:** Klik tombol dan tunggu AI bekerja.
    4. **Download:** Simpan hasil dalam format Word.
    """)
    st.divider()
    st.info("Aplikasi ini mendukung Kurikulum Merdeka & Standar HOTS.")

# --- KONTEN UTAMA ---
st.title("📝 Generator Soal Ujian Otomatis")
st.write("Ubah materi pelajaran menjadi Kisi-kisi, Soal, dan Kartu Soal dalam sekejap.")

col1, col2 = st.columns([2, 1])

with col1:
    st.markdown('<div class="main-card">', unsafe_allow_html=True)
    metode = st.radio("Sumber Materi:", ["Tempel Teks Manual", "Unggah Dokumen PDF"])
    
    materi_final = ""
    if metode == "Tempel Teks Manual":
        materi_final = st.text_area("Masukkan Materi:", height=250, placeholder="Contoh: Bab 1 Fotosintesis...")
    else:
        file_pdf = st.file_uploader("Upload PDF Materi", type=["pdf"])
        if file_pdf:
            reader = PdfReader(file_pdf)
            for page in reader.pages:
                materi_final += page.extract_text()
            st.success("PDF Berhasil dibaca!")
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.write("⚙️ **Setelan Ujian**")
    jenjang = st.selectbox("Jenjang", ["SD", "SMP", "SMA/SMK"])
    jml = st.number_input("Jumlah Soal", 1, 30, 5)
    
    if st.button("Generate Sekarang ✨", use_container_width=True):
        if materi_final:
            prog = st.progress(0)
            msg = st.empty()
            
            msg.text("🔍 Menganalisis materi...")
            prog.progress(25)
            
            prompt = f"Buatlah kisi-kisi tabel, {jml} soal PG HOTS {jenjang}, dan kartu soal dari materi ini: {materi_final[:6000]}"
            
            try:
                response = model.generate_content(prompt)
                prog.progress(75)
                msg.text("📝 Menyusun format dokumen...")
                
                st.markdown("---")
                st.markdown(response.text)
                
                # Tombol Download
                doc_file = buat_file_word(response.text)
                st.download_button("📥 Download Hasil (Word)", doc_file, "Ujian_Otomatis.docx", use_container_width=True)
                
                prog.progress(100)
                msg.text("✅ Selesai!")
            except Exception as e:
                st.error(f"Gagal: {e}")
        else:
            st.warning("Masukkan materi terlebih dahulu!")

# --- FOOTER ---
st.markdown(f"""
    <div class="footer">
        <p>© 2024 {NAMA_APLIKASI} | Powered by Gemini AI | Dikembangkan oleh {NAMA_PEMBUAT}</p>
    </div>
    """, unsafe_allow_html=True)
